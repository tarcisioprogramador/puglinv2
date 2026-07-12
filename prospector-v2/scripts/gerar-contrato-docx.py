#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera contrato .docx protegido (somente leitura com regiões editáveis).
Uso: python3 gerar-contrato-docx.py dados.json saida.docx

Requer: pip install python-docx

dados.json deve conter: NOME_CLIENTE, CPF_CNPJ_CLIENTE, ENDERECO_CLIENTE,
CIDADE_UF_CLIENTE, NOME_PRESTADOR, CPF_CNPJ_PRESTADOR, ENDERECO_PRESTADOR,
CIDADE_UF_PRESTADOR, URL_SITE_ANTIGO, URL_PUBLICADA, VALOR, VALOR_EXTENSO,
FORMA_PAGAMENTO, PRAZO_ENTREGA, RODADAS_AJUSTES,
MANUTENCAO (true/false), VALOR_MANUTENCAO, TEXTO_HOSPEDAGEM,
CIDADE_FORO, CIDADE_ASSINATURA
"""

import json, sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def main():
    if len(sys.argv) < 3:
        print('Uso: python3 gerar-contrato-docx.py dados.json saida.docx')
        sys.exit(1)

    d = json.load(open(sys.argv[1], encoding='utf-8'))
    pid = [100]
    doc = Document()

    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.2)
        s.left_margin = s.right_margin = Cm(2.2)

    def par(texto='', bold=False, center=False, size=11, antes=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(antes)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
        if texto:
            r = p.add_run(texto)
            r.bold = bold
            r.font.size = Pt(size)
            r.font.name = 'Georgia'
        return p

    def run(p, texto, bold=False, size=11):
        r = p.add_run(texto)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = 'Georgia'
        return r

    def editavel(p, texto):
        pid[0] += 1
        p_id = str(pid[0])
        ps = OxmlElement('w:permStart')
        ps.set(qn('w:id'), p_id)
        ps.set(qn('w:edGrp'), 'everyone')
        p._p.append(ps)
        r = run(p, texto)
        r.font.highlight_color = 7  # amarelo
        pe = OxmlElement('w:permEnd')
        pe.set(qn('w:id'), p_id)
        p._p.append(pe)

    def campo(p, valor, rotulo):
        if 'preencher' in (valor or '').lower() or not valor:
            editavel(p, f' [{rotulo}: preencher aqui] ')
        else:
            run(p, valor)

    par('CONTRATO DE PRESTACAO DE SERVICOS', bold=True, center=True, size=13)
    par('CRIACAO E PUBLICACAO DE PAGINA NA INTERNET', bold=True, center=True, size=11)

    # CONTRATANTE
    p = par()
    run(p, 'CONTRATANTE: ', bold=True)
    run(p, d['NOME_CLIENTE'] + ', ')
    label = 'CPF/CNPJ'
    run(p, f'{label} nº ')
    campo(p, d.get('CPF_CNPJ_CLIENTE'), label)
    run(p, ', com endereco em ')
    campo(p, d.get('ENDERECO_CLIENTE'), 'endereco')
    run(p, ', ' + d['CIDADE_UF_CLIENTE'] + '.')

    # CONTRATADO
    p = par()
    run(p, 'CONTRATADO(A): ', bold=True)
    run(p, f"{d['NOME_PRESTADOR']}, CPF/CNPJ nº {d['CPF_CNPJ_PRESTADOR']}, com endereco em {d['ENDERECO_PRESTADOR']}, {d['CIDADE_UF_PRESTADOR']}.")

    par('As partes acima identificadas celebram o presente contrato de prestacao de servicos, que se regera pelas clausulas seguintes.')

    def clausula(n, titulo, texto):
        par(f'Clausula {n}a - {titulo}', bold=True, antes=12)
        par(texto)

    clausula(1, 'Do objeto',
        f'O presente contrato tem por objeto a criacao de nova versao da pagina na internet do CONTRATANTE ({d["URL_SITE_ANTIGO"]}), incluindo: redesign completo do layout com manutencao da identidade visual, redacao aprimorada do conteudo existente, adaptacao para dispositivos moveis e publicacao em {d["URL_PUBLICADA"]}.')

    clausula(2, 'Do valor e forma de pagamento',
        f'Pelos servicos descritos na Clausula 1a, o CONTRATANTE pagara ao CONTRATADO(A) o valor total de R$ {d["VALOR"]} ({d["VALOR_EXTENSO"]}), na seguinte forma: {d["FORMA_PAGAMENTO"]}.')

    clausula(3, 'Do prazo de entrega',
        f'A pagina em sua versao final sera entregue e publicada em ate {d["PRAZO_ENTREGA"]} a contar da assinatura deste contrato. Esta incluida {d["RODADAS_AJUSTES"]} rodada(s) de ajustes.')

    n = 4
    if d.get('MANUTENCAO'):
        clausula(4, 'Da manutencao mensal',
            f'O CONTRATANTE contrata servico de manutencao mensal pelo valor de R$ {d["VALOR_MANUTENCAO"]} mensais, com renovacao automatica.')
        n = 5

    clausula(n, 'Do conteudo e responsabilidades',
        'O CONTRATANTE declara ser titular dos direitos de uso de todos os textos, imagens e informacoes fornecidos.')

    clausula(n+1, 'Da hospedagem e dominio', d.get('TEXTO_HOSPEDAGEM', 'A hospedagem sera realizada em servidor contratado pelo CONTRATADO(A).'))

    clausula(n+2, 'Da rescicao',
        'Este contrato podera ser rescindido por qualquer das partes mediante comunicacao por escrito. Em caso de rescicao apos o inicio dos trabalhos, sera devido o valor proporcional.')

    clausula(n+3, 'Do foro',
        f'Fica eleito o foro da comarca de {d["CIDADE_FORO"]} para dirimir controversias.')

    p = par(antes=18)
    run(p, d['CIDADE_ASSINATURA'] + ', ')
    editavel(p, ' [data] ')
    run(p, '.')

    par('', antes=24)
    p = par(antes=18)
    run(p, '__________________________________________')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = par(antes=0)
    run(p, d['NOME_CLIENTE'] + ' - Contratante', bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    editavel(p, ' [assine aqui] ')

    p = par(antes=18)
    run(p, '__________________________________________')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = par(antes=0, center=True)
    run(p, d['NOME_PRESTADOR'] + ' - Contratado(a)', bold=True)

    par('Este documento e uma minuta base gerada automaticamente. Recomenda-se revisao juridica. Gerado pelo Prospector de Sites.', size=8, antes=20)

    # Protecao do documento
    dp = OxmlElement('w:documentProtection')
    dp.set(qn('w:edit'), 'readOnly')
    dp.set(qn('w:enforcement'), '1')
    doc.settings.element.append(dp)

    doc.save(sys.argv[2])
    print(f'docx gerado: {sys.argv[2]}')

if __name__ == '__main__':
    main()
